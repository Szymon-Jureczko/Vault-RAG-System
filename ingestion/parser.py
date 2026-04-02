"""Multimodal document parsers — PyMuPDF, RapidOCR, python-docx.

Provides DOCX parsing (python-docx), fast PDF extraction (PyMuPDF),
and OCR for scanned images and PDFs (RapidOCR via ONNX Runtime).
Each parser returns a list of text chunks with source metadata.

Usage::

    from ingestion.parser import parse_file

    result = parse_file(Path("report.pdf"), chunk_size=2000)
    for chunk in result.chunks:
        print(chunk.text, chunk.metadata)
"""

from __future__ import annotations

import logging
import uuid
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any

from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)

# ── Supported extension sets ────────────────────────────────────────────────
PDF_EXTENSIONS = {".pdf"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".tiff", ".tif", ".bmp"}
TEXT_EXTENSIONS = {".txt", ".md", ".rst", ".csv", ".json", ".xml", ".html"}
OFFICE_EXTENSIONS = {".docx"}
XLSX_EXTENSIONS = {".xlsx"}
EML_EXTENSIONS = {".eml"}

# ── OCR gating constants ────────────────────────────────────────────────────
_OCR_MIN_FILE_BYTES: int = 10_000  # skip images < 10 KB
_OCR_VARIANCE_THRESHOLD: float = 100.0  # skip near-blank/solid-colour images
_OCR_TARGET_WIDTH: int = 1200  # rescale for OCR — lower = less memory

# ── PDF text-density probe ──────────────────────────────────────────────────
_PDF_TEXT_THRESHOLD: int = 50  # avg chars/page below this → scanned


class Chunk(BaseModel):
    """A single text chunk extracted from a document.

    Attributes:
        chunk_id: Unique identifier for this chunk.
        text: The extracted text content.
        metadata: Source metadata (filename, page, parser, etc.).
    """

    chunk_id: str
    text: str
    metadata: dict = Field(default_factory=dict)


class ParserResult(BaseModel):
    """Output from parsing a single document.

    Attributes:
        file_path: Path to the source document.
        chunks: Extracted text chunks.
        parser_name: Parser that produced the result.
        page_count: Number of pages in the source (if applicable).
        success: Whether parsing completed without errors.
        error: Error description if success is False.
    """

    file_path: str
    chunks: list[Chunk] = Field(default_factory=list)
    parser_name: str = ""
    page_count: int = 0
    success: bool = True
    error: str | None = None


def _chunk_id(source: Path, index: int) -> str:
    """Generate a unique chunk identifier."""
    return f"{source.stem}_{index}_{uuid.uuid4().hex[:8]}"


def _split_text(
    text: str,
    source: Path,
    chunk_size: int,
    parser_name: str,
    page: int | None = None,
    overlap: int = 150,
) -> list[Chunk]:
    """Split text into fixed-size chunks with metadata.

    Args:
        text: Full extracted text.
        source: Source file path.
        chunk_size: Target characters per chunk.
        parser_name: Name of the parser that produced the text.
        page: Optional page number for page-level chunks.
        overlap: Characters of overlap between consecutive chunks.

    Returns:
        List of Chunk instances.
    """
    text = text.strip()
    if not text:
        return []

    # Prefix with filename so document names are searchable
    prefix = f"[Source: {source.name}]\n"

    chunks: list[Chunk] = []
    stride = max(1, chunk_size - overlap)
    for i in range(0, len(text), stride):
        segment = prefix + text[i : i + chunk_size]
        meta: dict = {
            "source": str(source),
            "filename": source.name,
            "parser": parser_name,
            "chunk_index": len(chunks),
        }
        if page is not None:
            meta["page"] = page
        chunks.append(
            Chunk(
                chunk_id=_chunk_id(source, len(chunks)),
                text=segment,
                metadata=meta,
            )
        )
    return chunks


# ── Base parser ─────────────────────────────────────────────────────────────


class BaseParser(ABC):
    """Abstract base for document parsers."""

    @property
    @abstractmethod
    def name(self) -> str:
        """Human-readable parser name."""

    @property
    @abstractmethod
    def extensions(self) -> set[str]:
        """File extensions this parser handles (lowercase, with dot)."""

    @abstractmethod
    def parse(
        self, path: Path, chunk_size: int = 1000, chunk_overlap: int = 200
    ) -> ParserResult:
        """Parse a document into text chunks.

        Args:
            path: Path to the source file.
            chunk_size: Approximate characters per chunk.
            chunk_overlap: Characters of overlap between consecutive chunks.

        Returns:
            ParserResult with extracted chunks.
        """


# ── Docling parser (layout-aware, table → markdown) ────────────────────────


class DocxParser(BaseParser):
    """Lightweight DOCX parser using python-docx.

    Extracts paragraph text and tables as pipe-delimited rows.
    Much faster than DoclingParser (no ML model loading).
    """

    @property
    def name(self) -> str:
        return "docx"

    @property
    def extensions(self) -> set[str]:
        return OFFICE_EXTENSIONS

    def parse(
        self, path: Path, chunk_size: int = 1000, chunk_overlap: int = 150
    ) -> ParserResult:
        """Parse DOCX using python-docx.

        Args:
            path: Path to the document.
            chunk_size: Approximate characters per chunk.

        Returns:
            ParserResult with text chunks.
        """
        try:
            from docx import Document
        except ImportError as exc:
            return ParserResult(
                file_path=str(path),
                success=False,
                error=f"python-docx not installed: {exc}",
                parser_name=self.name,
            )

        try:
            doc = Document(str(path))
            parts: list[str] = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            for table in doc.tables:
                rows: list[str] = []
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    parts.append("\n".join(rows))

            full_text = "\n\n".join(parts)
            chunks = _split_text(
                full_text, path, chunk_size, self.name, overlap=chunk_overlap
            )
            return ParserResult(
                file_path=str(path),
                chunks=chunks,
                parser_name=self.name,
                page_count=1,
                success=True,
            )
        except Exception as exc:
            logger.error("DocxParser failed on %s: %s", path, exc)
            return ParserResult(
                file_path=str(path),
                success=False,
                error=str(exc),
                parser_name=self.name,
            )


# ── PyMuPDF parser (fast PDF) ──────────────────────────────────────────────


class PyMuPDFParser(BaseParser):
    """Fast PDF text extraction using PyMuPDF."""

    @property
    def name(self) -> str:
        return "pymupdf"

    @property
    def extensions(self) -> set[str]:
        return PDF_EXTENSIONS

    def parse(
        self, path: Path, chunk_size: int = 2000, chunk_overlap: int = 150
    ) -> ParserResult:
        """Extract text from PDF pages using PyMuPDF (fitz).

        Args:
            path: Path to the PDF file.
            chunk_size: Approximate characters per chunk (default 2000
                to keep most pages in a single chunk).

        Returns:
            ParserResult with page-level chunks.
        """
        try:
            import fitz
        except ImportError as exc:
            return ParserResult(
                file_path=str(path),
                success=False,
                error=f"PyMuPDF not installed: {exc}",
                parser_name=self.name,
            )

        try:
            all_chunks: list[Chunk] = []
            with fitz.open(str(path)) as doc:
                for i in range(len(doc)):
                    page = doc[i]
                    page_num = i + 1
                    text = str(page.get_text())
                    page_chunks = _split_text(
                        text, path, chunk_size, self.name, page=page_num, overlap=chunk_overlap
                    )
                    all_chunks.extend(page_chunks)
                page_count = len(doc)

            return ParserResult(
                file_path=str(path),
                chunks=all_chunks,
                parser_name=self.name,
                page_count=page_count,
                success=True,
            )
        except Exception as exc:
            logger.error("PyMuPDF failed on %s: %s", path, exc)
            return ParserResult(
                file_path=str(path),
                success=False,
                error=str(exc),
                parser_name=self.name,
            )


# ── Scanned PDF parser (PyMuPDF + RapidOCR) ─────────────────────────────────


class ScannedPDFParser(BaseParser):
    """Parser for scanned PDFs using PyMuPDF page rendering + RapidOCR.

    Renders each page to a greyscale image via PyMuPDF, applies the same
    variance gate and preprocessing as ``RapidOCRParser``, then runs
    RapidOCR directly.  Much lighter than Docling (3 ONNX models instead
    of 5) and processes page-by-page to limit peak memory.
    """

    @property
    def name(self) -> str:
        return "scanned_pdf"

    @property
    def extensions(self) -> set[str]:
        return PDF_EXTENSIONS

    def parse(
        self, path: Path, chunk_size: int = 2000, chunk_overlap: int = 150
    ) -> ParserResult:
        """Extract text from a scanned PDF via page-level OCR.

        Args:
            path: Path to the PDF file.
            chunk_size: Approximate characters per chunk (default 2000
                to keep most pages in a single chunk).

        Returns:
            ParserResult with OCR-extracted chunks.
        """
        try:
            import fitz
            import numpy as np
        except ImportError as exc:
            return ParserResult(
                file_path=str(path),
                success=False,
                error=f"PyMuPDF/numpy not installed: {exc}",
                parser_name=self.name,
            )

        try:
            engine = RapidOCRParser._get_engine()
            all_chunks: list[Chunk] = []
            with fitz.open(str(path)) as doc:
                for i in range(len(doc)):
                    page = doc[i]
                    page_num = i + 1
                    pix = page.get_pixmap(dpi=150, colorspace=fitz.csGRAY)
                    gray = (
                        np.frombuffer(
                            pix.samples,
                            dtype=np.uint8,
                        )
                        .reshape(pix.height, pix.width)
                        .copy()
                    )
                    del pix  # free pixmap immediately

                    # Skip blank / near-blank pages
                    if float(gray.var()) < _OCR_VARIANCE_THRESHOLD:
                        del gray
                        continue

                    # Rescale if needed
                    from PIL import Image, ImageOps

                    img = Image.fromarray(gray)
                    del gray  # free numpy array
                    img = ImageOps.autocontrast(img)
                    w, h = img.size
                    if w < 800 or w > 3000:
                        scale = _OCR_TARGET_WIDTH / w
                        img = img.resize(
                            (int(w * scale), int(h * scale)),
                            Image.Resampling.LANCZOS,
                        )

                    ocr_arr = np.array(img)
                    del img  # free PIL image before inference
                    result, _ = engine(ocr_arr)
                    del ocr_arr  # free numpy array
                    text = "\n".join(line[1] for line in result) if result else ""
                    page_chunks = _split_text(
                        text,
                        path,
                        chunk_size,
                        self.name,
                        page=page_num,
                        overlap=chunk_overlap,
                    )
                    all_chunks.extend(page_chunks)

                page_count = len(doc)

            return ParserResult(
                file_path=str(path),
                chunks=all_chunks,
                parser_name=self.name,
                page_count=page_count,
                success=True,
            )
        except Exception as exc:
            logger.error("ScannedPDFParser failed on %s: %s", path, exc)
            return ParserResult(
                file_path=str(path),
                success=False,
                error=str(exc),
                parser_name=self.name,
            )


# ── RapidOCR parser (images, replaces Tesseract) ──────────────────────────


class RapidOCRParser(BaseParser):
    """OCR parser for scanned images using RapidOCR (ONNX Runtime).

    Replaces TesseractParser with in-process inference (no subprocess
    overhead) and two fast gates to skip non-text images.
    """

    _engine: Any | None = None  # class-level lazy singleton per process

    @classmethod
    def _get_engine(cls):
        if cls._engine is None:
            from rapidocr_onnxruntime import RapidOCR

            cls._engine = RapidOCR(det_limit_side_len=1280)
        return cls._engine

    @property
    def name(self) -> str:
        return "rapidocr"

    @property
    def extensions(self) -> set[str]:
        return IMAGE_EXTENSIONS

    def parse(
        self, path: Path, chunk_size: int = 1000, chunk_overlap: int = 150
    ) -> ParserResult:
        """Extract text from an image using RapidOCR.

        Args:
            path: Path to the image file.
            chunk_size: Approximate characters per chunk.

        Returns:
            ParserResult with OCR-extracted chunks.
        """
        try:
            import numpy as np
            from PIL import Image
        except ImportError as exc:
            return ParserResult(
                file_path=str(path),
                success=False,
                error=f"numpy/Pillow not installed: {exc}",
                parser_name=self.name,
            )

        try:
            # Gate 1: skip tiny files (icons, bullets, decorative)
            if path.stat().st_size < _OCR_MIN_FILE_BYTES:
                logger.debug("Skipping OCR on undersized image %s", path)
                return ParserResult(
                    file_path=str(path),
                    chunks=[],
                    parser_name=self.name,
                    page_count=1,
                    success=True,
                )

            image = Image.open(path)

            # Gate 2: pixel-variance — blank/solid-colour images
            gray_arr = np.array(image.convert("L"))
            if float(gray_arr.var()) < _OCR_VARIANCE_THRESHOLD:
                logger.debug("OCR skipped: low-variance image %s", path)
                return ParserResult(
                    file_path=str(path),
                    chunks=[],
                    parser_name=self.name,
                    page_count=1,
                    success=True,
                )

            # Preprocessing: greyscale + autocontrast + rescale to optimal OCR width
            gray = image.convert("L")
            from PIL import ImageOps

            gray = ImageOps.autocontrast(gray)
            del image, gray_arr  # free original image + variance array
            w, h = gray.size
            if w < 800 or w > 3000:
                scale = _OCR_TARGET_WIDTH / w
                gray = gray.resize(
                    (int(w * scale), int(h * scale)),
                    Image.Resampling.LANCZOS,
                )

            ocr_arr = np.array(gray)
            del gray  # free PIL image before OCR inference
            result, _ = self._get_engine()(ocr_arr)
            del ocr_arr  # free numpy array immediately
            text = "\n".join(line[1] for line in result) if result else ""
            chunks = _split_text(text, path, chunk_size, self.name, page=1, overlap=chunk_overlap)

            return ParserResult(
                file_path=str(path),
                chunks=chunks,
                parser_name=self.name,
                page_count=1,
                success=True,
            )
        except Exception as exc:
            logger.error("RapidOCR failed on %s: %s", path, exc)
            return ParserResult(
                file_path=str(path),
                success=False,
                error=str(exc),
                parser_name=self.name,
            )


# ── Plain-text parser ──────────────────────────────────────────────────────


class TextParser(BaseParser):
    """Parser for plain-text and structured-text files."""

    @property
    def name(self) -> str:
        return "text"

    @property
    def extensions(self) -> set[str]:
        return TEXT_EXTENSIONS

    def parse(self, path: Path, chunk_size: int = 1000, chunk_overlap: int = 150) -> ParserResult:
        """Read and chunk plain-text files.

        Args:
            path: Path to the text file.
            chunk_size: Approximate characters per chunk.

        Returns:
            ParserResult with text chunks.
        """
        try:
            text = path.read_text(encoding="utf-8", errors="replace")
            chunks = _split_text(text, path, chunk_size, self.name, overlap=chunk_overlap)

            return ParserResult(
                file_path=str(path),
                chunks=chunks,
                parser_name=self.name,
                page_count=1,
                success=True,
            )
        except Exception as exc:
            logger.error("TextParser failed on %s: %s", path, exc)
            return ParserResult(
                file_path=str(path),
                success=False,
                error=str(exc),
                parser_name=self.name,
            )


# ── XLSX parser (openpyxl, lightweight) ───────────────────────────────────


class OpenpyxlParser(BaseParser):
    """Lightweight XLSX parser using openpyxl instead of Docling."""

    @property
    def name(self) -> str:
        return "openpyxl"

    @property
    def extensions(self) -> set[str]:
        return XLSX_EXTENSIONS

    def parse(self, path: Path, chunk_size: int = 1000, chunk_overlap: int = 150) -> ParserResult:
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            return ParserResult(
                file_path=str(path),
                success=False,
                error=f"openpyxl not installed: {exc}",
                parser_name=self.name,
            )

        try:
            wb = load_workbook(str(path), read_only=True, data_only=True)
            parts: list[str] = []
            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                rows: list[str] = [f"## Sheet: {sheet_name}"]
                for row in ws.iter_rows(values_only=True):
                    cells = [str(c) if c is not None else "" for c in row]
                    rows.append(" | ".join(cells))
                parts.append("\n".join(rows))
            wb.close()

            full_text = "\n\n".join(parts)
            chunks = _split_text(full_text, path, chunk_size, self.name, overlap=chunk_overlap)
            return ParserResult(
                file_path=str(path),
                chunks=chunks,
                parser_name=self.name,
                page_count=len(wb.sheetnames) if parts else 0,
                success=True,
            )
        except Exception as exc:
            logger.error("OpenpyxlParser failed on %s: %s", path, exc)
            return ParserResult(
                file_path=str(path),
                success=False,
                error=str(exc),
                parser_name=self.name,
            )


# ── EML parser ─────────────────────────────────────────────────────────────


class EmlParser(BaseParser):
    """Parser for RFC-822 email files (.eml) using stdlib ``email``."""

    @property
    def name(self) -> str:
        return "eml"

    @property
    def extensions(self) -> set[str]:
        return EML_EXTENSIONS

    def parse(self, path: Path, chunk_size: int = 1000, chunk_overlap: int = 150) -> ParserResult:
        """Extract headers and body text from an EML file.

        Args:
            path: Path to the .eml file.
            chunk_size: Approximate characters per chunk.

        Returns:
            ParserResult with email content as text chunks.
        """
        import email as _email

        try:
            msg = _email.message_from_bytes(path.read_bytes())
            parts: list[str] = []

            # Include key headers as context
            for header in ("Subject", "From", "To", "CC", "Date"):
                val = msg.get(header)
                if val:
                    parts.append(f"{header}: {val}")

            # Extract all text/plain and text/html payloads
            for part in msg.walk():
                ct = part.get_content_type()
                if ct not in ("text/plain", "text/html"):
                    continue
                charset = part.get_content_charset() or "utf-8"
                payload = part.get_payload(decode=True)
                if isinstance(payload, bytes):
                    parts.append(payload.decode(charset, errors="replace"))

            full_text = "\n\n".join(parts)
            chunks = _split_text(full_text, path, chunk_size, self.name, overlap=chunk_overlap)
            return ParserResult(
                file_path=str(path),
                chunks=chunks,
                parser_name=self.name,
                page_count=1,
                success=True,
            )
        except Exception as exc:
            logger.error("EmlParser failed on %s: %s", path, exc)
            return ParserResult(
                file_path=str(path),
                success=False,
                error=str(exc),
                parser_name=self.name,
            )


# ── PDF text-density probe ──────────────────────────────────────────────────


def is_scanned_pdf(file_path: Path) -> bool:
    """Probe PDF text density to decide whether OCR is needed.

    Opens the file, extracts raw text from all pages, and returns True
    if the average character count per page is below
    ``_PDF_TEXT_THRESHOLD``.  The probe costs ~5-20 ms versus 1-30 s
    for a full OCR pass.

    Args:
        file_path: Path to the PDF.

    Returns:
        True if the PDF is likely scanned / requires OCR.
    """
    try:
        import fitz

        with fitz.open(str(file_path)) as doc:
            total_chars = sum(len(str(page.get_text()).strip()) for page in doc)
            avg = total_chars / max(len(doc), 1)
        return avg < _PDF_TEXT_THRESHOLD
    except Exception:
        return True  # assume scanned — OCR attempt is safer than losing content


# ── Dispatcher ──────────────────────────────────────────────────────────────

# Non-PDF parsers — PDFs are routed dynamically by parse_file().
_PARSERS: list[BaseParser] = [
    OpenpyxlParser(),
    DocxParser(),
    RapidOCRParser(),
    TextParser(),
    EmlParser(),
]

_EXT_MAP: dict[str, BaseParser] = {}
for _p in _PARSERS:
    for _ext in _p.extensions:
        if _ext not in _EXT_MAP:
            _EXT_MAP[_ext] = _p

SUPPORTED_EXTENSIONS = set(_EXT_MAP.keys()) | PDF_EXTENSIONS


def parse_file(
    path: Path,
    chunk_size: int = 2000,
    chunk_overlap: int = 150,
) -> ParserResult:
    """Parse a file using the best available parser.

    PDFs are routed via a cheap text-density probe:
    - Text-layer PDFs  -> ``PyMuPDFParser`` (fast, no OCR)
    - Scanned PDFs     -> ``ScannedPDFParser`` (PyMuPDF + RapidOCR)

    All other files are routed by extension via ``_EXT_MAP``.

    Args:
        path: Path to the source document.
        chunk_size: Approximate characters per chunk.
        chunk_overlap: Characters of overlap between consecutive chunks.

    Returns:
        ParserResult from the matched parser.
    """
    ext = path.suffix.lower()

    if ext in PDF_EXTENSIONS:
        if is_scanned_pdf(path):
            parser: BaseParser = ScannedPDFParser()
        else:
            parser = PyMuPDFParser()
    else:
        parser = _EXT_MAP.get(ext)  # type: ignore[assignment]

    if parser is None:
        return ParserResult(
            file_path=str(path),
            success=False,
            error=f"Unsupported extension: {ext}",
        )
    logger.info("Parsing %s with %s", path.name, parser.name)
    return parser.parse(path, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
