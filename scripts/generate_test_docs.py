"""Generate 1000 synthetic test documents across PDF, DOCX, XLSX, EML, and
scanned image formats with varied, realistic content using Faker.

Run from the workspace root::

    python scripts/generate_test_docs.py

Output: data/test_docs/  (200 of each format)
"""

from __future__ import annotations

import email.mime.multipart
import email.mime.text
import random
import textwrap
from pathlib import Path

from docx import Document
from faker import Faker
from fpdf import FPDF
from openpyxl import Workbook
from PIL import Image, ImageDraw, ImageFont

fake = Faker()
OUT_DIR = Path("data/test_docs")
COUNT_PER_FORMAT = 200  # 5 formats × 200 = 1000 total

# ── Content templates ────────────────────────────────────────────────────────

TOPICS = [
    "quarterly financial report",
    "employee performance review",
    "project status update",
    "vendor contract summary",
    "technical specification",
    "meeting minutes",
    "risk assessment",
    "compliance audit",
    "product roadmap",
    "incident report",
    "research findings",
    "budget proposal",
    "onboarding guide",
    "security policy",
    "market analysis",
]


def _body_paragraphs(n: int = 4) -> list[str]:
    """Return n paragraphs of fake body text."""
    topic = random.choice(TOPICS)
    intro = (
        f"This document concerns the {topic} for {fake.company()}. "
        f"Prepared by {fake.name()} on {fake.date_this_decade()}."
    )
    paragraphs = [intro] + [fake.paragraph(nb_sentences=6) for _ in range(n - 1)]
    return paragraphs


def _title() -> str:
    return f"{random.choice(TOPICS).title()} — {fake.company()}"


# ── PDF ──────────────────────────────────────────────────────────────────────


def generate_pdf(path: Path) -> None:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", style="B", size=16)
    pdf.multi_cell(0, 10, _title())
    pdf.ln(4)
    pdf.set_font("Helvetica", size=11)
    for para in _body_paragraphs(5):
        pdf.multi_cell(0, 7, para)
        pdf.ln(3)
    # Optional table
    pdf.ln(4)
    pdf.set_font("Helvetica", style="B", size=11)
    pdf.cell(60, 8, "Item", border=1)
    pdf.cell(40, 8, "Q1", border=1)
    pdf.cell(40, 8, "Q2", border=1)
    pdf.ln()
    pdf.set_font("Helvetica", size=10)
    for _ in range(4):
        pdf.cell(60, 7, fake.bs(), border=1)
        pdf.cell(40, 7, str(fake.random_int(100, 9999)), border=1)
        pdf.cell(40, 7, str(fake.random_int(100, 9999)), border=1)
        pdf.ln()
    pdf.output(str(path))


# ── DOCX ─────────────────────────────────────────────────────────────────────


def generate_docx(path: Path) -> None:
    doc = Document()
    doc.add_heading(_title(), level=1)
    doc.add_paragraph(f"Author: {fake.name()}   |   Date: {fake.date_this_decade()}")
    doc.add_paragraph("")
    for i, para in enumerate(_body_paragraphs(5), start=1):
        doc.add_heading(f"Section {i}", level=2)
        doc.add_paragraph(para)
    # Table
    doc.add_heading("Summary Table", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text, hdr[1].text, hdr[2].text = "Metric", "Value", "Status"
    for _ in range(5):
        row = table.add_row().cells
        row[0].text = fake.bs().title()
        row[1].text = str(fake.random_int(1, 500))
        row[2].text = random.choice(["On Track", "At Risk", "Completed", "Pending"])
    doc.save(str(path))


# ── XLSX ─────────────────────────────────────────────────────────────────────

SHEET_TYPES = ["Financial", "Inventory", "HR", "Sales", "Operations"]


def generate_xlsx(path: Path) -> None:
    wb = Workbook()
    ws = wb.active
    sheet_type = random.choice(SHEET_TYPES)
    ws.title = sheet_type

    if sheet_type == "Financial":
        headers = ["Date", "Description", "Category", "Debit", "Credit", "Balance"]
        ws.append(headers)
        balance = fake.random_int(10000, 100000)
        for _ in range(random.randint(30, 60)):
            debit = random.choice([fake.random_int(0, 5000), 0])
            credit = fake.random_int(0, 3000) if debit == 0 else 0
            balance += credit - debit
            ws.append(
                [
                    str(fake.date_this_year()),
                    fake.bs(),
                    random.choice(
                        ["Payroll", "Operations", "Marketing", "IT", "Travel"]
                    ),
                    debit or "",
                    credit or "",
                    balance,
                ]
            )
    elif sheet_type == "Inventory":
        headers = ["SKU", "Product", "Category", "Qty", "Unit Price", "Total Value"]
        ws.append(headers)
        for _ in range(random.randint(40, 80)):
            qty = fake.random_int(1, 500)
            price = round(random.uniform(1.5, 999.99), 2)
            ws.append(
                [
                    fake.bothify("SKU-####-??").upper(),
                    fake.catch_phrase(),
                    fake.bs().split()[0].title(),
                    qty,
                    price,
                    round(qty * price, 2),
                ]
            )
    elif sheet_type == "HR":
        headers = ["Employee ID", "Name", "Department", "Role", "Start Date", "Salary"]
        ws.append(headers)
        for _ in range(random.randint(20, 50)):
            ws.append(
                [
                    fake.bothify("EMP-#####"),
                    fake.name(),
                    random.choice(
                        ["Engineering", "Finance", "HR", "Marketing", "Legal"]
                    ),
                    fake.job(),
                    str(fake.date_this_decade()),
                    fake.random_int(40000, 180000),
                ]
            )
    else:
        headers = ["ID", "Name", "Value", "Region", "Date", "Status"]
        ws.append(headers)
        for _ in range(random.randint(30, 70)):
            ws.append(
                [
                    fake.uuid4()[:8].upper(),
                    fake.company(),
                    fake.random_int(100, 50000),
                    fake.country(),
                    str(fake.date_this_year()),
                    random.choice(["Active", "Closed", "Pending"]),
                ]
            )

    # Second summary sheet
    ws2 = wb.create_sheet("Summary")
    ws2.append(["Report", _title()])
    ws2.append(["Generated", str(fake.date_this_decade())])
    ws2.append(["Prepared by", fake.name()])
    ws2.append([])
    for para in _body_paragraphs(2):
        ws2.append([para])

    wb.save(str(path))


# ── EML ──────────────────────────────────────────────────────────────────────


def generate_eml(path: Path) -> None:
    msg = email.mime.multipart.MIMEMultipart("alternative")
    msg["Subject"] = f"Re: {_title()}"
    msg["From"] = fake.email()
    msg["To"] = ", ".join(fake.email() for _ in range(random.randint(1, 3)))
    msg["CC"] = fake.email()
    msg["Date"] = fake.date_time_this_year().strftime("%a, %d %b %Y %H:%M:%S +0000")
    msg["Message-ID"] = f"<{fake.uuid4()}@{fake.domain_name()}>"

    body_lines = [
        f"Hi {fake.first_name()},",
        "",
        *[textwrap.fill(p, width=80) for p in _body_paragraphs(3)],
        "",
        f"Please review the attached {random.choice(TOPICS)} at your earliest convenience.",
        "",
        "Best regards,",
        fake.name(),
        fake.job(),
        fake.company(),
        fake.phone_number(),
    ]
    body = "\n".join(body_lines)
    msg.attach(email.mime.text.MIMEText(body, "plain"))

    # HTML version
    html = f"<html><body><p>{'</p><p>'.join(body_lines)}</p></body></html>"
    msg.attach(email.mime.text.MIMEText(html, "html"))

    path.write_text(msg.as_string(), encoding="utf-8")


# ── Scanned Image (PNG with text, for OCR) ───────────────────────────────────


def generate_scanned_image(path: Path) -> None:
    """Create a white A4-ish image with black printed text simulating a scan."""
    img = Image.new("RGB", (1240, 1754), color=(245, 245, 245))  # ~A4 at 150 dpi
    draw = ImageDraw.Draw(img)

    try:
        font_title = ImageFont.truetype(
            "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 28
        )
        font_body = ImageFont.truetype(
            "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 20
        )
    except OSError:
        font_title = ImageFont.load_default()
        font_body = ImageFont.load_default()

    y = 80
    # Title
    draw.text((80, y), _title(), fill=(10, 10, 10), font=font_title)
    y += 60
    draw.text(
        (80, y),
        f"Author: {fake.name()}   |   {fake.date_this_decade()}",
        fill=(60, 60, 60),
        font=font_body,
    )
    y += 50

    for para in _body_paragraphs(6):
        # Wrap text to fit image width
        wrapped = textwrap.wrap(para, width=85)
        for line in wrapped:
            if y > 1650:
                break
            draw.text((80, y), line, fill=(20, 20, 20), font=font_body)
            y += 28
        y += 14

    # Add slight noise to simulate scan
    import random as _rnd

    for _ in range(800):
        x, yy = _rnd.randint(0, 1239), _rnd.randint(0, 1753)
        gray = _rnd.randint(180, 220)
        img.putpixel((x, yy), (gray, gray, gray))

    img.save(str(path), format="PNG", optimize=True)


# ── Main ─────────────────────────────────────────────────────────────────────


def main() -> None:
    generators = {
        "pdf": (generate_pdf, ".pdf"),
        "docx": (generate_docx, ".docx"),
        "xlsx": (generate_xlsx, ".xlsx"),
        "eml": (generate_eml, ".eml"),
        "img": (generate_scanned_image, ".png"),
    }

    for fmt, (fn, ext) in generators.items():
        fmt_dir = OUT_DIR / fmt
        fmt_dir.mkdir(parents=True, exist_ok=True)
        print(f"Generating {COUNT_PER_FORMAT} {fmt.upper()} files...", flush=True)
        for i in range(COUNT_PER_FORMAT):
            out_path = fmt_dir / f"{fmt}_{i:04d}{ext}"
            try:
                fn(out_path)
            except Exception as exc:  # noqa: BLE001
                print(f"  [WARN] {out_path.name}: {exc}")
            if (i + 1) % 50 == 0:
                print(f"  {i + 1}/{COUNT_PER_FORMAT} done", flush=True)

    total = sum(1 for _ in OUT_DIR.rglob("*") if _.is_file())
    print(f"\nDone. {total} files written to {OUT_DIR}/")


if __name__ == "__main__":
    main()
